# app/utils/resume_extractor.py
"""
Text extraction utilities for PDF and DOCX resume files.
"""
import PyPDF2
from docx import Document
from typing import Optional
import logging
from io import BytesIO

logger = logging.getLogger(__name__)


def extract_text_from_pdf(file_bytes: BytesIO) -> Optional[str]:
    """
    Extract text from PDF file bytes.
    
    Args:
        file_bytes: BytesIO object containing PDF file data
        
    Returns:
        Extracted text string or None if extraction fails
    """
    try:
        pdf_reader = PyPDF2.PdfReader(file_bytes)
        text_parts = []
        
        for page in pdf_reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        
        full_text = "\n".join(text_parts).strip()
        
        if not full_text:
            logger.warning("PDF text extraction returned empty string")
            return None
            
        logger.info(f"Successfully extracted {len(full_text)} characters from PDF")
        return full_text
        
    except Exception as e:
        logger.error(f"PDF extraction failed: {str(e)}", exc_info=True)
        return None


def extract_text_from_docx(file_bytes: BytesIO) -> Optional[str]:
    """
    Extract text from DOCX file bytes.
    
    Args:
        file_bytes: BytesIO object containing DOCX file data
        
    Returns:
        Extracted text string or None if extraction fails
    """
    try:
        doc = Document(file_bytes)
        text_parts = []
        
        # Extract from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():
                text_parts.append(para.text)
        
        # Extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        text_parts.append(cell.text)
        
        full_text = "\n".join(text_parts).strip()
        
        if not full_text:
            logger.warning("DOCX text extraction returned empty string")
            return None
            
        logger.info(f"Successfully extracted {len(full_text)} characters from DOCX")
        return full_text
        
    except Exception as e:
        logger.error(f"DOCX extraction failed: {str(e)}", exc_info=True)
        return None


def extract_text(file_bytes: BytesIO, file_type: str) -> Optional[str]:
    """
    Extract text from resume file based on type.
    
    Args:
        file_bytes: BytesIO object containing file data
        file_type: File extension ('pdf' or 'docx')
        
    Returns:
        Extracted text string or None if extraction fails
    """
    file_type = file_type.lower().replace('.', '')
    
    if file_type == 'pdf':
        return extract_text_from_pdf(file_bytes)
    elif file_type == 'docx':
        return extract_text_from_docx(file_bytes)
    else:
        logger.error(f"Unsupported file type: {file_type}")
        return None
